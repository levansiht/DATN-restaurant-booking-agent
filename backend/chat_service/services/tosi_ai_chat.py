import os
import json
import time
from django.http import StreamingHttpResponse
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from api_chat_bot import settings
from langchain_community.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from common.services.llm_service import get_llm_service, LLMProvider
import docx
from ..models import create_document_embedding, Chat, Message
from ..serializers import ChatHistoryListSerializer, ChatHistoryDetailSerializer

class TosiAiChatService:
    def __init__(self, llm_provider: LLMProvider = LLMProvider.OPENAI):
        self.llm_provider = llm_provider
        self.llm_service = get_llm_service()
        
        # Get API key based on provider
        if llm_provider == LLMProvider.OPENAI:
            api_key = settings.OPENAI_API_KEY
            if not api_key:
                raise ValueError("OPENAI_API_KEY environment variable is not set")
        elif llm_provider == LLMProvider.CLAUDE:
            api_key = os.getenv("ANTHROPIC_API_KEY")
            if not api_key:
                raise ValueError("ANTHROPIC_API_KEY environment variable is not set")
        else:
            api_key = None

        model = "gpt-4o-mini" if llm_provider == LLMProvider.OPENAI else "claude-3-sonnet-20240229"

        self.llm = self.llm_service.create_agent_llm(
            provider=llm_provider,
            model=model,
            temperature=0.7,
            streaming=True,
            api_key=api_key,
            verbose=False,
        )
        
        # Configurable streaming delay (in seconds)
        self.streaming_delay = 0.05  # 50ms default delay
    def get_chat_history(self, user):
        chats = Chat.objects.filter(user=user, is_deleted=False)
        return ChatHistoryListSerializer(chats, many=True).data
    
    def get_chat_history_detail(self, chat_id):
        chat = Chat.objects.get(uuid=chat_id, is_deleted=False)
        return ChatHistoryDetailSerializer(chat).data
        
    def chat(self, request, data):
        user = request.user
        chat_id = data.get('chat_id', None)
        message = data.get('message', None)
        chat = self.get_chat_by_id(user, chat_id, message)
        history = self.get_history_by_chat_id(chat_id)

        return self.stream_chat(data, chat, history) 

    def get_chat_by_id(self, user, chat_id, title=None):
        if not chat_id:
            chat = Chat.objects.create(user=user, title=title if title and len(title) <= 50 else title[:50])
            return chat
        chat = Chat.objects.get(user=user, uuid=chat_id, is_deleted=False)
        return chat
    
    def get_history_by_chat_id(self, chat_id):
        messages = Message.objects.filter(chat__uuid=chat_id).order_by("created_at")
        # Build history input for LLM as a list of dicts with role and content
        history = []
        for msg in messages:
            if msg.sender == Message.Sender.HUMAN:
                history.append({"role": "user", "content": msg.message})
            else:
                history.append({"role": "assistant", "content": msg.message})
        return history

    def stream_chat(self, input_data, chat, history):
        """
        Stream chat response using LangChain and OpenAI with whitepaper context.
        
        Args:
            input_data: Dictionary containing user input, typically {'message': 'user message'}
        """
        user_message = input_data.get('message', '')
        
        if not user_message:
            return self._create_error_response('No message provided')

        def event_stream():
            try:
                # Get relevant context and build messages
                context_content = self._get_context_content(user_message)
                system_message = self._build_system_message(context_content)
                messages = self._build_messages(system_message, history, user_message)
                
                # Stream the response
                output_message = ""
                for chunk in self.llm.stream(messages):
                    if hasattr(chunk, 'content') and chunk.content:
                        output_message += chunk.content
                        yield self._format_stream_data('token', content=chunk.content)
                        # Add delay to control streaming speed (adjust value as needed)
                        time.sleep(self.streaming_delay)  # Use configurable delay
                
                # Send end signal and save messages
                yield self._format_stream_data('end')
                self._save_conversation_messages(chat, user_message, output_message)
                
            except Exception as e:
                error_message = self._handle_streaming_error(e)
                yield self._format_stream_data('error', error=error_message)
                yield self._format_stream_data('end')

        response = StreamingHttpResponse(event_stream(), content_type='text/event-stream')
        response['Cache-Control'] = 'no-cache'
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Headers'] = 'Cache-Control'        
        return response

    def _get_context_content(self, user_message: str) -> str:
        """Get relevant context content from vector database."""
        relevant_docs = self.search_whitepaper_embeddings(user_message, top_k=3)
        if relevant_docs:
            return "\n\n".join(doc.page_content for doc in relevant_docs)
        return ""

    def _build_system_message(self, context_content: str) -> str:
        """Build the system message with context."""
        return f"""You are a helpful AI assistant with access to a whitepaper about Tosi Growth Holding. 
        Use the following relevant whitepaper content to answer questions accurately and comprehensively:

        {context_content}

        When answering questions:
        1. Base your responses on the whitepaper content provided above
        2. If the question is not related to the whitepaper, politely redirect to whitepaper-related topics
        3. Provide specific information from the whitepaper when possible
        4. Be helpful and informative while staying within the context of the whitepaper
        5. If you don't have enough information to answer the question, say so and suggest what information might be available
        """

    def _build_messages(self, system_message: str, history: list, user_message: str) -> list:
        """Build the complete message list for the LLM."""
        messages = [SystemMessage(content=system_message)]
        
        # Add conversation history
        for hist_msg in history:
            if hist_msg["role"] == "user":
                messages.append(HumanMessage(content=hist_msg["content"]))
            elif hist_msg["role"] == "assistant":
                messages.append(AIMessage(content=hist_msg["content"]))
        
        # Add current user message
        messages.append(HumanMessage(content=user_message))
        return messages

    def _format_stream_data(self, data_type: str, **kwargs) -> str:
        """Format streaming data as Server-Sent Events."""
        data = {'type': data_type}
        data.update(kwargs)
        return f"data: {json.dumps(data)}\n\n"

    def _handle_streaming_error(self, error: Exception) -> str:
        """Handle and format streaming errors."""
        error_message = str(error)
        
        if "429" in error_message or "Too Many Requests" in error_message:
            return "Rate limit exceeded. Please wait a moment and try again."
        elif "401" in error_message or "Unauthorized" in error_message:
            return "Invalid API key. Please check your OpenAI API key."
        elif "400" in error_message:
            return "Invalid request. Please check your input."
        elif "500" in error_message:
            return "OpenAI service is temporarily unavailable. Please try again later."
        else:
            return f"Error generating response: {error_message}"

    def _create_error_response(self, error_message: str):
        """Create an error response for invalid requests."""
        def error_stream():
            yield self._format_stream_data('error', error=error_message)
        
        response = StreamingHttpResponse(error_stream(), content_type='text/event-stream')
        response['Cache-Control'] = 'no-cache'
        return response

    def _save_conversation_messages(self, chat, user_message: str, bot_message: str):
        """Save both user and bot messages to the database."""
        self.save_message(chat, user_message, Message.Sender.HUMAN)
        self.save_message(chat, bot_message, Message.Sender.BOT)
    
    def save_message(self, chat, message, sender):
        Message.objects.create(chat=chat, message=message, sender=sender)
    

    def _create_documents_from_text(self):
        """
        Read the Whitepaper Tosi Growth Holding.docx file and store it in DocumentEmbedding as vector database.
        This method processes the DOCX file, splits it into chunks, and creates embeddings for each chunk.
        """
        try:
            # Path to the whitepaper file
            whitepaper_path = "Whitepaper Tosi Growth Holding.docx"
            
            if not os.path.exists(whitepaper_path):
                print(f"Error: Whitepaper file not found at {whitepaper_path}")
                return
            
            print(f"Reading whitepaper from: {whitepaper_path}")
            
            # Read the DOCX file
            doc = docx.Document(whitepaper_path)
            
            # Extract text from all paragraphs
            full_text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    full_text += paragraph.text.strip() + "\n\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text.strip() + "\n"
                    full_text += "\n"
            
            if not full_text.strip():
                print("Warning: No text content found in the whitepaper file")
                return
            
            print(f"Extracted {len(full_text)} characters from whitepaper")
            
            # Create a Document object
            document = Document(
                page_content=full_text,
                metadata={
                    "source": "Whitepaper Tosi Growth Holding.docx",
                    "document_type": "whitepaper",
                    "company": "Tosi Growth Holding",
                    "extraction_method": "python-docx"
                }
            )
            
            # Split the document into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=300,
                chunk_overlap=100,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            documents = text_splitter.split_documents([document])
            print(f"Split document into {len(documents)} chunks")
            
            # Store documents in vector database
            create_document_embedding(documents)
            
            print("Successfully processed and stored whitepaper in vector database")
            
            # Also save the full content to JSON for backward compatibility
            self._save_whitepaper_to_json(full_text)
            
        except Exception as e:
            print(f"Error processing whitepaper: {e}")
            raise

    def _save_whitepaper_to_json(self, content):
        """
        Save the whitepaper content to JSON file for backward compatibility.
        """
        try:
            data = {
                "content": content,
                "source": "Whitepaper Tosi Growth Holding.docx",
                "processed_at": time.strftime("%Y-%m-%d %H:%M:%S")
            }
            
            with open('whitepaper_data.json', 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            print("Saved whitepaper content to whitepaper_data.json")
            
        except Exception as e:
            print(f"Error saving whitepaper to JSON: {e}")

    def search_whitepaper_embeddings(self, query: str, top_k: int = 5):
        """
        Search the whitepaper embeddings for relevant content based on a query.
        
        Args:
            query: The search query
            top_k: Number of top results to return
            
        Returns:
            List of relevant document chunks
        """
        try:
            from ..models import get_pgvector_client, DocumentEmbedding
            
            # Check if we have any embeddings in the database
            if DocumentEmbedding.objects.count() == 0:
                print("No embeddings found in database. Using fallback content.")
                return []
            
            # Get the vector store client
            vector_store = get_pgvector_client()
            
            # Search for similar documents
            results = vector_store.similarity_search(query, k=top_k)

            print("================================================")
            print(results)
            print("================================================")
            
            return results
            
        except Exception as e:
            print(f"Error searching whitepaper embeddings: {e}")
            # Return empty list to fall back to full content
            return []
